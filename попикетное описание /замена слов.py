import openpyxl
from docx import Document
import tkinter as tk
from tkinter import filedialog, font

def process_files():
    docx_file_path = docx_file.get()
    xlsx_file_path = xlsx_file.get()

    workbook = openpyxl.load_workbook(xlsx_file_path)
    sheet = workbook.active

    doc = Document(docx_file_path)

    new_doc = Document()

    for row in range(1, 110):  
        values = [sheet.cell(row, i+1).value for i in range(len(word_entries))]

        if all(value is not None for value in values):
            for paragraph in doc.paragraphs:
                updated_text = paragraph.text
                for word_entry, value in zip(word_entries, values):
                    updated_text = updated_text.replace(word_entry.get(), str(value))

                new_doc.add_paragraph(updated_text)

    new_doc.save('new_text.docx')

def add_field():
    word_label = tk.Label(root, text="Cлово для замены:", font=large_font)
    word_label.grid(row=len(word_entries)+3, column=0)

    word_entry = tk.Entry(root, font=large_font, width=8)
    word_entry.grid(row=len(word_entries)+3, column=1)

    word_entries.append(word_entry)

def select_file(var, button):
    file_path = filedialog.askopenfilename()
    if file_path:
        var.set(file_path)
        button.config(background='green')

root = tk.Tk()
root.title("Обработчик файлов")

large_font = font.Font(size=30)

docx_file = tk.StringVar()
xlsx_file = tk.StringVar()

word_entries = []

docx_label = tk.Label(root, text="Файл ворда", font=large_font)
docx_label.grid(row=0, column=0)

docx_button = tk.Button(root, text="Выбрать", command=lambda: select_file(docx_file, docx_button), font=large_font)
docx_button.grid(row=0, column=1)

xlsx_label = tk.Label(root, text="Файл экселя", font=large_font)
xlsx_label.grid(row=1, column=0)

xlsx_button = tk.Button(root, text="Выбрать", command=lambda: select_file(xlsx_file, xlsx_button), font=large_font)
xlsx_button.grid(row=1, column=1)

add_field_button = tk.Button(root, text="+", command=add_field, font=large_font)
add_field_button.grid(row=99, column=0, columnspan=2)

for _ in range(3):
    add_field()

process_button = tk.Button(root, text="Обработать", command=process_files, font=large_font)
process_button.grid(row=100, column=0, columnspan=2)

root.mainloop()
