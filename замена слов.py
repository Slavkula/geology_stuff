import openpyxl
from docx import Document

# Открываем файл Excel
workbook = openpyxl.load_workbook('1.xlsx')
sheet = workbook['Лист1']

# Открываем файл Word
doc = Document('1.docx')

# Создаем новый файл Word
new_doc = Document()

# Цикл для обработки строк в Excel
for row in range(1, 110):  # Предполагая, что вам нужно обработать 100 строк
    # Получаем значения из ячеек
    value1 = sheet.cell(row, 1).value
    value2 = sheet.cell(row, 2).value
    value3 = sheet.cell(row, 3).value
    value4 = sheet.cell(row, 4).value
    value5 = sheet.cell(row, 5).value
    value6 = sheet.cell(row, 6).value
    value7 = sheet.cell(row, 7).value

    # Проверяем, что значения не пустые
    if value1 is not None and value2 is not None and value3 is not None:
        # Обновляем значения в тексте
        for paragraph in doc.paragraphs:
            updated_text = paragraph.text.replace('ЭТО1', str(value1))
            updated_text = updated_text.replace('ЭТО2', str(value2))
            updated_text = updated_text.replace('ЭТО3', str(value3))
            updated_text = updated_text.replace('ЭТО4', str(value4))
            updated_text = updated_text.replace('ЭТО5', str(value5))
            updated_text = updated_text.replace('ЭТО6', str(value6))
            updated_text = updated_text.replace('ЭТО7', str(value7))
            new_doc.add_paragraph(updated_text)

# Сохраняем новый файл Word
new_doc.save('new_text.docx')
