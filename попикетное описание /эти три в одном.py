from docx import Document
import re

def remove_zeroes_v1(doc_path, save_path):
    doc = Document(doc_path)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = re.sub(r'(\d+,\d)0\b', r'\1', run.text)
    doc.save(save_path)

def remove_zeroes_v2(save_path):
    doc = Document(save_path)
    for para in doc.paragraphs:
        for run in para.runs:
            for i in range(0, 131):
                num = i / 10
                run.text = re.sub(rf'({num:.2f})м\b', lambda m: f'{m.group(1)[:-1]}м', run.text)
    doc.save(save_path)

def replace_text_in_word_doc(save_path):
    doc = Document(save_path)
    pattern = r"(встречены в пределах) скважин(: \d+[.;])"
    replacement = r"\1 скважины\2"

    for paragraph in doc.paragraphs:
        if re.search(pattern, paragraph.text):
            paragraph.text = re.sub(pattern, replacement, paragraph.text)

    doc.save(save_path)

# Выполнение скриптов по очереди
doc_path = 'попикетное описание 3 участок.docx'
save_path = 'готовое.docx'
remove_zeroes_v1(doc_path, save_path)
remove_zeroes_v2(save_path)
replace_text_in_word_doc(save_path)
